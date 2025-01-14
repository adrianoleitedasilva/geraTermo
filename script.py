import pandas as pd
from docx import Document
import os

def criar_arquivo(nome, cpf, equipamento, patrimonio):
    try:
        doc = Document()
        doc.add_heading('TERMO DE ENTREGA DE EQUIPAMENTO', level=1)

        doc.add_paragraph(
            f"\n\n\nEu, {nome}, portador(a) do CPF nº {cpf}, confirmo o recebimento do(s) equipamento(s) descrito(s) abaixo, de propriedade da empresa, sob os seguintes termos e condições:")
        
        doc.add_heading('Equipamento(s) Entregue(s):', level=3)
        doc.add_paragraph(f"\n- Equipamento: {equipamento}")
        doc.add_paragraph(f"- Número de Patrimônio: {patrimonio}")

        doc.add_paragraph("\nDeclaro estar ciente de que:")
        doc.add_paragraph("1. O(s) equipamento(s) recebido(s) é(são) de propriedade exclusiva da empresa e deve(m) ser utilizado(s) exclusivamente para fins profissionais relacionados às atividades da mesma.")
        doc.add_paragraph("2. Estou ciente de que deverei zelar pela conservação e bom uso do(s) equipamento(s), mantendo-o(s) em boas condições de funcionamento.")
        doc.add_paragraph("3. Caso o(s) equipamento(s) seja(m) danificado(s) ou extraviado(s) por negligência, descuido ou uso inadequado, poderei ser responsabilizado(a) pelos custos de reparo ou substituição.")
        doc.add_paragraph("4. Deverei devolver o(s) equipamento(s) nas mesmas condições em que foram entregues, salvo desgaste natural pelo uso, quando solicitado pela empresa ou no término do vínculo empregatício.")

        doc.add_heading('\nAssinatura e Data:', level=3)
        doc.add_paragraph("Assinatura do(a) colaborador(a): _________________________________________")
        doc.add_paragraph("Data: ____/____/________")

        doc.add_paragraph("\nResponsável pela Entrega:")
        doc.add_paragraph("Nome: _____________________________________________________________")
        doc.add_paragraph("Cargo: ____________________________________________________________")

        file_name = f"{nome}_{cpf}.docx"
        doc.save(file_name)

        return "OK"
    except Exception as e:
        print(f"Erro ao criar arquivo para {nome}: {e}")
        return "Não OK"

def processar_planilha(caminho_entrada, caminho_saida):
    try:
        # Ler a planilha
        df = pd.read_excel(caminho_entrada)

        # Verificar se as colunas necessárias existem
        colunas_necessarias = ["Nome", "CPF", "Equipamento", "Patrimônio"]
        for coluna in colunas_necessarias:
            if coluna not in df.columns:
                raise Exception(f"Coluna obrigatória '{coluna}' ausente na planilha.")

        # Criar uma nova coluna para status
        df["Status"] = ""

        # Processar cada linha
        for index, row in df.iterrows():
            nome = row["Nome"]
            cpf = row["CPF"]
            equipamento = row["Equipamento"]
            patrimonio = row["Patrimônio"]

            status = criar_arquivo(nome, cpf, equipamento, patrimonio)
            df.at[index, "Status"] = status

        # Salvar a planilha atualizada
        df.to_excel(caminho_saida, index=False)
        print("Processamento concluído. Arquivos gerados e planilha atualizada.")

    except Exception as e:
        print(f"Erro ao processar a planilha: {e}")

# Exemplo de uso:
caminho_entrada = "planilha_equipamentos.xlsx"  # Substitua pelo caminho da sua planilha
caminho_saida = "planilha_equipamentos_atualizada.xlsx"  # Nome para a planilha atualizada
processar_planilha(caminho_entrada, caminho_saida)
